from __future__ import annotations

import io
import re
import unicodedata
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any

import pandas as pd
from dateutil import parser as dateparser
from docx import Document

FRENCH_MONTHS = {
    "janvier": 1, "fevrier": 2, "février": 2, "mars": 3, "avril": 4, "mai": 5, "juin": 6,
    "juillet": 7, "aout": 8, "août": 8, "septembre": 9, "octobre": 10, "novembre": 11, "decembre": 12, "décembre": 12,
}

ALIASES = {
    "title": ["activite", "activité", "activity", "tache", "tâche", "task", "action", "resultat", "résultat", "livrable"],
    "description": ["description", "details", "détails", "objectif", "observations", "commentaire", "justification", "indicateur"],
    "start_date": ["date debut", "date début", "debut", "début", "start", "date depart", "date départ", "periode debut", "période début"],
    "end_date": ["date fin", "fin", "deadline", "echeance", "échéance", "due", "date retour", "periode fin", "période fin"],
    "planned_date": ["jour planifie", "jour planifié", "date planifiee", "date planifiée", "planning", "jour", "date"],
    "month": ["mois", "month", "periode", "période"],
    "project": ["projet", "project", "programme", "bailleur", "donor"],
    "partner": ["partenaire", "partner", "collaborateur", "bailleur", "donateur", "financement"],
    "responsible": ["responsable", "owner", "lead", "chef", "personne responsable", "qui"],
    "priority": ["priorite", "priorité", "urgence", "urgent", "priority"],
    "status": ["statut", "status", "etat", "état", "avancement"],
    "budget_estimated": ["budget", "montant", "cout", "coût", "estimation", "ariary", "ar"],
    "location": ["lieu", "site", "localisation", "commune", "village", "terrain"],
}

PROJECT_KEYWORDS = {
    "SOS Lemurs": ["sos lemur", "sos lemurs", "sos lémur", "sos lémurs", "lemur", "lémur"],
    "Darwin Initiatives": ["darwin", "darwin initiatives", "darwin initiative"],
    "Seacology": ["seacology"],
    "Rainforest Trust": ["rainforest", "rainforest trust", "rft"],
}

PARTNER_KEYWORDS = {
    "TGBS (MBG)": ["tgbs", "mbg", "missouri botanical"],
    "MfM": ["mfm"],
    "UWE": ["uwe"],
    "Regen": ["regen"],
    "UNI": ["uni", "universite", "université"],
    "ENS": ["ens", "ecole normale", "école normale"],
}


def strip_accents(text: str) -> str:
    text = str(text or "")
    return "".join(c for c in unicodedata.normalize("NFD", text) if unicodedata.category(c) != "Mn")


def norm(text: str) -> str:
    text = strip_accents(text).lower().strip()
    text = re.sub(r"[^a-z0-9]+", " ", text)
    return re.sub(r"\s+", " ", text).strip()


def find_column(columns: list[str], logical_name: str) -> str | None:
    normalized = {c: norm(c) for c in columns}
    for c, nc in normalized.items():
        for alias in ALIASES.get(logical_name, []):
            na = norm(alias)
            if nc == na or na in nc or nc in na:
                return c
    return None


def parse_date(value: Any) -> str | None:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return None
    if isinstance(value, (pd.Timestamp, datetime)):
        return value.date().isoformat()
    text = str(value).strip()
    if not text or text.lower() in {"nan", "nat", "none"}:
        return None
    # Excel serial if it arrived as a number in text form
    try:
        if re.fullmatch(r"\d{5}(\.0)?", text):
            dt = pd.to_datetime(float(text), unit="D", origin="1899-12-30")
            return dt.date().isoformat()
    except Exception:
        pass
    # Month only => first day of current/mentioned year
    low = text.lower()
    for mname, mnum in FRENCH_MONTHS.items():
        if mname in low:
            year_match = re.search(r"(20\d{2}|19\d{2})", low)
            year = int(year_match.group(1)) if year_match else datetime.today().year
            return datetime(year, mnum, 1).date().isoformat()
    try:
        return dateparser.parse(text, dayfirst=True, fuzzy=True).date().isoformat()
    except Exception:
        return None


def parse_amount(value: Any) -> float:
    if value is None or (isinstance(value, float) and pd.isna(value)):
        return 0.0
    if isinstance(value, (int, float)):
        return float(value)
    text = strip_accents(str(value)).lower().replace("ariary", "").replace("ar", "")
    text = re.sub(r"[^0-9,\.\-]", "", text)
    if text.count(",") == 1 and text.count(".") == 0:
        text = text.replace(",", ".")
    else:
        text = text.replace(",", "")
    try:
        return float(text)
    except Exception:
        return 0.0


def detect_name(text: str, lookup: dict[str, int], keyword_map: dict[str, list[str]] | None = None) -> tuple[str | None, int | None]:
    if not text:
        return None, None
    ntext = norm(text)
    tokens = set(ntext.split())
    for name, _id in lookup.items():
        nn = norm(name)
        if not nn:
            continue
        if len(nn) <= 3:
            matched = nn in tokens
        else:
            matched = nn in ntext or ntext in nn
        if matched:
            return name, _id
    if keyword_map:
        tokens = set(ntext.split())
        for name, keys in keyword_map.items():
            for key in keys:
                nk = norm(key)
                if not nk:
                    continue
                if len(nk) <= 3:
                    matched = nk in tokens
                else:
                    matched = nk in ntext
                if matched:
                    # Try exact current lookup name
                    for opt, _id in lookup.items():
                        if norm(opt) == norm(name):
                            return opt, _id
                    return name, None
    return None, None


def detect_priority(text: str) -> str:
    n = norm(text)
    if any(k in n for k in ["critique", "tres urgent", "très urgent", "immédiat", "immediat"]):
        return "Critique"
    if any(k in n for k in ["urgent", "haute", "eleve", "élevé", "important"]):
        return "Haute"
    if any(k in n for k in ["faible", "low"]):
        return "Faible"
    return "Normale"


def detect_status(text: str) -> str:
    n = norm(text)
    if any(k in n for k in ["termine", "terminé", "done", "acheve", "réalisé", "realise"]):
        return "Terminée"
    if any(k in n for k in ["en cours", "ongoing", "progress"]):
        return "En cours"
    if any(k in n for k in ["reporte", "reporté", "postpone"]):
        return "Reportée"
    if any(k in n for k in ["annule", "annulé", "cancel"]):
        return "Annulée"
    if any(k in n for k in ["planifie", "planifié", "scheduled"]):
        return "Planifiée"
    return "À planifier"


def smart_score(row: dict[str, Any]) -> tuple[float, str]:
    checks = {
        "titre": bool(row.get("title")),
        "description": bool(row.get("description")),
        "date": bool(row.get("planned_date") or row.get("start_date") or row.get("end_date")),
        "projet": bool(row.get("project_id")),
        "responsable": bool(row.get("responsible_member_id") or row.get("responsible_text")),
    }
    score = sum(checks.values()) / len(checks) * 100
    missing = [k for k, ok in checks.items() if not ok]
    notes = "Complet" if not missing else "À compléter : " + ", ".join(missing)
    return round(score, 1), notes


def map_dataframe(df: pd.DataFrame, projects: dict[str, int], partners: dict[str, int], members: dict[str, int], source_file: str, source_sheet: str = "") -> list[dict[str, Any]]:
    if df.empty:
        return []
    # Drop empty rows/cols
    df = df.dropna(how="all").dropna(axis=1, how="all")
    df.columns = [str(c).strip() if str(c).strip() else f"Colonne_{i+1}" for i, c in enumerate(df.columns)]
    cols = list(df.columns)
    mapping = {k: find_column(cols, k) for k in ALIASES}
    rows = []
    for idx, raw in df.iterrows():
        joined = " | ".join(str(v) for v in raw.tolist() if str(v).strip() and str(v).lower() != "nan")
        if not joined.strip():
            continue
        title_col = mapping.get("title")
        desc_col = mapping.get("description")
        title = str(raw.get(title_col, "")).strip() if title_col else ""
        desc = str(raw.get(desc_col, "")).strip() if desc_col else ""
        if not title:
            # Try to use the first text-rich cell as title
            text_cells = [str(v).strip() for v in raw.tolist() if isinstance(v, str) and len(str(v).strip()) > 2]
            if text_cells:
                title = text_cells[0][:180]
        if not title and not desc:
            continue
        if not desc:
            desc_parts = []
            for c in cols:
                if c != title_col and str(raw.get(c, "")).strip() and str(raw.get(c, "")).lower() != "nan":
                    desc_parts.append(f"{c}: {raw.get(c)}")
            desc = "; ".join(desc_parts)[:1500]
        project_text = str(raw.get(mapping.get("project"), "")) if mapping.get("project") else joined
        partner_text = str(raw.get(mapping.get("partner"), "")) if mapping.get("partner") else joined
        resp_text = str(raw.get(mapping.get("responsible"), "")) if mapping.get("responsible") else ""
        project_name, project_id = detect_name(project_text + " " + title + " " + desc, projects, PROJECT_KEYWORDS)
        partner_name, partner_id = detect_name(partner_text + " " + title + " " + desc, partners, PARTNER_KEYWORDS)
        resp_name, resp_id = detect_name(resp_text, members, None)
        start_date = parse_date(raw.get(mapping.get("start_date"))) if mapping.get("start_date") else None
        end_date = parse_date(raw.get(mapping.get("end_date"))) if mapping.get("end_date") else None
        planned_date = parse_date(raw.get(mapping.get("planned_date"))) if mapping.get("planned_date") else None
        if not planned_date and mapping.get("month"):
            planned_date = parse_date(raw.get(mapping.get("month")))
        priority_text = str(raw.get(mapping.get("priority"), "")) if mapping.get("priority") else joined
        status_text = str(raw.get(mapping.get("status"), "")) if mapping.get("status") else joined
        loc = str(raw.get(mapping.get("location"), "")).strip() if mapping.get("location") else ""
        budget = parse_amount(raw.get(mapping.get("budget_estimated"))) if mapping.get("budget_estimated") else 0.0
        out = {
            "title": title[:240] or "Activité importée",
            "description": desc,
            "activity_type": "Activité projet",
            "project_id": project_id,
            "project_detected": project_name or "",
            "partner_id": partner_id,
            "partner_detected": partner_name or "",
            "responsible_member_id": resp_id,
            "responsible_text": resp_text,
            "location": loc,
            "start_date": start_date,
            "end_date": end_date,
            "planned_date": planned_date or start_date or end_date,
            "status": detect_status(status_text),
            "priority": detect_priority(priority_text),
            "budget_estimated": budget,
            "source_file": source_file,
            "source_sheet": source_sheet,
            "source_row": str(idx + 2),
        }
        out["smart_score"], out["smart_notes"] = smart_score(out)
        rows.append(out)
    return rows


def read_excel_smart(file_bytes: bytes, filename: str, projects: dict[str, int], partners: dict[str, int], members: dict[str, int]) -> list[dict[str, Any]]:
    all_rows: list[dict[str, Any]] = []
    xls = pd.ExcelFile(io.BytesIO(file_bytes))
    for sheet in xls.sheet_names:
        # Try normal header first, then scan for better header row if necessary.
        raw = pd.read_excel(xls, sheet_name=sheet, header=None)
        if raw.dropna(how="all").empty:
            continue
        header_row = 0
        best_score = -1
        for i in range(min(12, len(raw))):
            vals = [str(v) for v in raw.iloc[i].tolist()]
            score = sum(1 for logical in ALIASES for a in ALIASES[logical] if any(norm(a) in norm(v) for v in vals))
            if score > best_score:
                best_score = score
                header_row = i
        df = pd.read_excel(xls, sheet_name=sheet, header=header_row)
        all_rows.extend(map_dataframe(df, projects, partners, members, filename, sheet))
    return all_rows


def read_csv_smart(file_bytes: bytes, filename: str, projects: dict[str, int], partners: dict[str, int], members: dict[str, int]) -> list[dict[str, Any]]:
    for sep in [",", ";", "\t"]:
        try:
            df = pd.read_csv(io.BytesIO(file_bytes), sep=sep)
            if df.shape[1] > 1:
                return map_dataframe(df, projects, partners, members, filename, "CSV")
        except Exception:
            continue
    return []


def table_to_df(table) -> pd.DataFrame:
    data = [[cell.text.strip() for cell in row.cells] for row in table.rows]
    if not data:
        return pd.DataFrame()
    header = data[0]
    if len(set(header)) != len(header):
        header = [f"Colonne_{i+1}" if not h else h for i, h in enumerate(header)]
    return pd.DataFrame(data[1:], columns=header)


def read_docx_smart(file_bytes: bytes, filename: str, projects: dict[str, int], partners: dict[str, int], members: dict[str, int]) -> list[dict[str, Any]]:
    doc = Document(io.BytesIO(file_bytes))
    rows: list[dict[str, Any]] = []
    for ti, table in enumerate(doc.tables):
        df = table_to_df(table)
        rows.extend(map_dataframe(df, projects, partners, members, filename, f"Table {ti+1}"))
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    blocks: list[str] = []
    current: list[str] = []
    starts = re.compile(r"^(activité|activite|action|t[âa]che|task|mission|résultat|resultat|\d+[\.)-])", re.I)
    for p in paragraphs:
        if starts.search(p) and current:
            blocks.append("\n".join(current))
            current = [p]
        else:
            current.append(p)
    if current:
        blocks.append("\n".join(current))
    if not blocks and paragraphs:
        blocks = paragraphs
    for i, block in enumerate(blocks):
        if len(block) < 15:
            continue
        lines = [l.strip(" -•\t") for l in block.splitlines() if l.strip()]
        title = lines[0][:240]
        desc = "\n".join(lines[1:]) if len(lines) > 1 else block
        joined = block
        project_name, project_id = detect_name(joined, projects, PROJECT_KEYWORDS)
        partner_name, partner_id = detect_name(joined, partners, PARTNER_KEYWORDS)
        resp_match = re.search(r"responsable\s*[:\-]\s*([^\n;]+)", joined, flags=re.I)
        resp_text = resp_match.group(1).strip() if resp_match else ""
        resp_name, resp_id = detect_name(resp_text, members, None)
        date_candidates = re.findall(r"\b\d{1,2}[/-]\d{1,2}[/-]\d{2,4}\b|\b(?:janvier|février|fevrier|mars|avril|mai|juin|juillet|août|aout|septembre|octobre|novembre|décembre|decembre)\s+20\d{2}\b", joined, flags=re.I)
        planned = parse_date(date_candidates[0]) if date_candidates else None
        end = parse_date(date_candidates[-1]) if len(date_candidates) > 1 else planned
        budget_match = re.search(r"(?:budget|montant|co[ûu]t)\s*[:\-]?\s*([0-9 .,'AriaryarAR]+)", joined, flags=re.I)
        out = {
            "title": title,
            "description": desc,
            "activity_type": "Activité projet",
            "project_id": project_id,
            "project_detected": project_name or "",
            "partner_id": partner_id,
            "partner_detected": partner_name or "",
            "responsible_member_id": resp_id,
            "responsible_text": resp_text,
            "location": "",
            "start_date": planned,
            "end_date": end,
            "planned_date": planned,
            "status": detect_status(joined),
            "priority": detect_priority(joined),
            "budget_estimated": parse_amount(budget_match.group(1)) if budget_match else 0,
            "source_file": filename,
            "source_sheet": "Paragraphes Word",
            "source_row": str(i + 1),
        }
        out["smart_score"], out["smart_notes"] = smart_score(out)
        rows.append(out)
    # de-duplicate simple duplicate titles from tables + paragraphs
    unique = []
    seen = set()
    for r in rows:
        key = norm(r.get("title", ""))[:80]
        if key and key not in seen:
            unique.append(r)
            seen.add(key)
    return unique


def analyze_file(file_bytes: bytes, filename: str, projects: dict[str, int], partners: dict[str, int], members: dict[str, int]) -> tuple[list[dict[str, Any]], dict[str, Any]]:
    suffix = Path(filename).suffix.lower()
    if suffix in [".xlsx", ".xls"]:
        rows = read_excel_smart(file_bytes, filename, projects, partners, members)
    elif suffix == ".csv":
        rows = read_csv_smart(file_bytes, filename, projects, partners, members)
    elif suffix == ".docx":
        rows = read_docx_smart(file_bytes, filename, projects, partners, members)
    else:
        rows = []
    summary = {
        "total_detected": len(rows),
        "missing_description": sum(1 for r in rows if not r.get("description")),
        "missing_date": sum(1 for r in rows if not (r.get("planned_date") or r.get("start_date") or r.get("end_date"))),
        "missing_project": sum(1 for r in rows if not r.get("project_id")),
        "missing_responsible": sum(1 for r in rows if not (r.get("responsible_member_id") or r.get("responsible_text"))),
        "urgent_or_high": sum(1 for r in rows if r.get("priority") in ["Haute", "Critique"]),
        "avg_smart_score": round(sum(r.get("smart_score", 0) for r in rows) / len(rows), 1) if rows else 0,
    }
    return rows, summary
