from __future__ import annotations

import io
import tempfile
from copy import copy
from datetime import date, datetime
from pathlib import Path
from typing import Any

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Cm, Pt
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Border, Font, PatternFill, Side
from openpyxl.utils import get_column_letter

APP_DIR = Path(__file__).resolve().parent.parent
LOGO_PATH = APP_DIR / "assets" / "logo_sadabe.png"
BUDGET_TEMPLATE = APP_DIR / "templates" / "rapport_financier_canevas.xlsx"

GREEN = "2E7D32"
LIGHT_GREEN = "E8F5E9"
DARK = "1B5E20"
BORDER = Side(style="thin", color="999999")


def _set_cell(ws, coord: str, value: Any) -> None:
    cell = ws[coord]
    # If coord is inside a merged range but is not the anchor, write to the anchor cell.
    if cell.__class__.__name__ == "MergedCell":
        for rng in ws.merged_cells.ranges:
            if coord in rng:
                ws.cell(rng.min_row, rng.min_col, value)
                return
    cell.value = value


def _parse_date(d: Any) -> date | None:
    if d in (None, "", pd.NaT):
        return None
    if isinstance(d, date):
        return d
    try:
        return pd.to_datetime(d).date()
    except Exception:
        return None


def _safe_name(text: str) -> str:
    return "".join(c for c in str(text) if c.isalnum() or c in " _-").strip().replace(" ", "_")[:50] or "document"


def activities_to_xlsx(df: pd.DataFrame) -> bytes:
    output = io.BytesIO()
    export_df = df.copy()
    with pd.ExcelWriter(output, engine="openpyxl") as writer:
        export_df.to_excel(writer, sheet_name="Planning_filtré", index=False)
        ws = writer.book["Planning_filtré"]
        ws.freeze_panes = "A2"
        for cell in ws[1]:
            cell.fill = PatternFill("solid", fgColor=GREEN)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
        for col in ws.columns:
            max_len = 10
            col_letter = get_column_letter(col[0].column)
            for cell in col:
                value = str(cell.value or "")
                max_len = min(max(max_len, len(value) + 2), 55)
                cell.alignment = Alignment(vertical="top", wrap_text=True)
            ws.column_dimensions[col_letter].width = max_len
        for row in ws.iter_rows(min_row=2):
            for cell in row:
                cell.border = Border(top=BORDER, left=BORDER, right=BORDER, bottom=BORDER)
    return output.getvalue()


def make_budget_workbook(request: dict[str, Any], lines: list[dict[str, Any]]) -> bytes:
    if BUDGET_TEMPLATE.exists():
        wb = load_workbook(BUDGET_TEMPLATE)
        ws = wb[wb.sheetnames[0]]
        ws.title = "Demande_budget"
        # Remove other personal sample tabs; keep Parametres_ONG if available.
        for s in list(wb.sheetnames):
            if s not in {"Demande_budget", "Parametres_ONG"}:
                del wb[s]
    else:
        wb = Workbook()
        ws = wb.active
        ws.title = "Demande_budget"
        _build_blank_budget_sheet(ws)

    # Header fields, based on the SADABE financial canevas structure.
    fields = {
        "C4": "SADABE",
        "C5": request.get("title", ""),
        "J5": request.get("responsible_name", ""),
        "C6": request.get("location", ""),
        "H6": request.get("project_name", ""),
        "C7": _parse_date(request.get("date_depart")),
        "F7": _parse_date(request.get("date_retour")),
        "H7": request.get("nb_days", 0),
        "C8": request.get("beneficiary_name", ""),
        "H8": date.today(),
    }
    for cell, value in fields.items():
        _set_cell(ws, cell, value)
    for c in ["C7", "F7", "H8"]:
        ws[c].number_format = "dd/mm/yyyy"

    # Clear and fill line items around the standard rows 12:18.
    start_row = 12
    max_lines = max(7, len(lines))
    if len(lines) > 7:
        ws.insert_rows(start_row + 7, len(lines) - 7)
        for r in range(start_row + 7, start_row + len(lines)):
            for col in range(1, 11):
                ws.cell(r, col)._style = copy(ws.cell(start_row + 6, col)._style)
    for i in range(max_lines):
        r = start_row + i
        line = lines[i] if i < len(lines) else {}
        ws.cell(r, 1, i + 1)
        ws.cell(r, 2, line.get("rubrique", ""))
        ws.cell(r, 3, line.get("base_calcul", ""))
        ws.cell(r, 4, float(line.get("qty") or 0))
        ws.cell(r, 5, float(line.get("unit_rate") or 0))
        # If the user provided a custom total and no qty/rate, keep it; otherwise formula.
        custom_total = float(line.get("total") or 0)
        if custom_total and not (line.get("qty") and line.get("unit_rate")):
            ws.cell(r, 6, custom_total)
        else:
            ws.cell(r, 6, f'=IF(AND(D{r}<>"",E{r}<>""),D{r}*E{r},0)')
        ws.cell(r, 7, float(line.get("received_before") or 0))
        ws.cell(r, 8, float(line.get("received_during") or 0))
        ws.cell(r, 9, f"=F{r}-G{r}-H{r}")
        ws.cell(r, 10, line.get("comment", ""))
        for col in range(1, 11):
            ws.cell(r, col).alignment = Alignment(wrap_text=True, vertical="top")
            ws.cell(r, col).border = Border(top=BORDER, left=BORDER, right=BORDER, bottom=BORDER)
    total_row = start_row + max_lines
    # If original total row is now farther, create/refresh it.
    ws.cell(total_row, 1, "TOTAL GÉNÉRAL")
    ws.cell(total_row, 6, f"=SUM(F{start_row}:F{total_row-1})")
    ws.cell(total_row, 7, f"=SUM(G{start_row}:G{total_row-1})")
    ws.cell(total_row, 8, f"=SUM(H{start_row}:H{total_row-1})")
    ws.cell(total_row, 9, f"=SUM(I{start_row}:I{total_row-1})")
    ws.cell(total_row, 10, "Le reste = Total à percevoir - Reçu avant mission - Reçu pendant mission")
    for col in range(1, 11):
        ws.cell(total_row, col).font = Font(bold=True)
        ws.cell(total_row, col).fill = PatternFill("solid", fgColor=LIGHT_GREEN)
        ws.cell(total_row, col).border = Border(top=BORDER, left=BORDER, right=BORDER, bottom=BORDER)

    # Synthesis block: robust placement after total.
    synth_row = total_row + 2
    ws.cell(synth_row, 1, "2. SYNTHÈSE FINANCIÈRE")
    ws.cell(synth_row, 1).font = Font(bold=True, color=DARK)
    ws.cell(synth_row + 1, 1, "Total des besoins de mission")
    ws.cell(synth_row + 1, 4, f"=F{total_row}")
    ws.cell(synth_row + 2, 1, "Total reçu avant mission")
    ws.cell(synth_row + 2, 4, f"=G{total_row}")
    ws.cell(synth_row + 3, 1, "Total reçu pendant mission")
    ws.cell(synth_row + 3, 4, f"=H{total_row}")
    ws.cell(synth_row + 4, 1, "Somme restant à percevoir")
    ws.cell(synth_row + 4, 4, f"=I{total_row}")
    for row in range(synth_row + 1, synth_row + 5):
        ws.cell(row, 4).number_format = '#,##0 "Ar"'

    # Number formats and widths.
    for row in ws.iter_rows(min_row=start_row, max_row=total_row, min_col=4, max_col=9):
        for cell in row:
            cell.number_format = '#,##0'
    widths = {"A": 8, "B": 24, "C": 28, "D": 12, "E": 14, "F": 16, "G": 16, "H": 18, "I": 16, "J": 36}
    for col, width in widths.items():
        ws.column_dimensions[col].width = width

    out = io.BytesIO()
    wb.save(out)
    return out.getvalue()


def _build_blank_budget_sheet(ws) -> None:
    ws.merge_cells("A1:J1")
    ws["A1"] = "DEMANDE DE BUDGET / AVANCE DE MISSION - SADABE"
    ws["A1"].font = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill = PatternFill("solid", fgColor=GREEN)
    ws["A1"].alignment = Alignment(horizontal="center")
    labels = [
        ("A4", "ONG"), ("F4", "Réf. mission"),
        ("A5", "Intitulé de la mission"), ("H5", "Responsable"),
        ("A6", "Lieu(x)"), ("F6", "Service / Projet"),
        ("A7", "Date départ"), ("D7", "Date retour"), ("G7", "Nb jours"),
        ("A8", "Bénéficiaire"), ("F8", "Date rapport"),
    ]
    for cell, value in labels:
        ws[cell] = value
        ws[cell].font = Font(bold=True)
    ws["A10"] = "1. DÉTAIL DES DROITS ET BESOINS DE MISSION"
    ws["A10"].font = Font(bold=True, color=DARK)
    headers = ["N°", "Rubrique", "Base de calcul", "Qté / Jours", "Taux unitaire", "Total à percevoir", "Reçu avant mission", "Reçu pendant mission", "Reste", "Commentaires / justification"]
    for c, h in enumerate(headers, 1):
        cell = ws.cell(11, c, h)
        cell.fill = PatternFill("solid", fgColor=GREEN)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)


def make_leave_docx(data: dict[str, Any]) -> bytes:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.6)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    header = section.header
    htbl = header.add_table(rows=1, cols=2, width=Cm(17))
    htbl.cell(0, 0).width = Cm(4)
    if LOGO_PATH.exists():
        p = htbl.cell(0, 0).paragraphs[0]
        p.add_run().add_picture(str(LOGO_PATH), width=Cm(2.7))
    htxt = htbl.cell(0, 1).paragraphs[0]
    htxt.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = htxt.add_run("ONG SADABE\nServices d’Appui pour le Développement Autonome\nde la Biodiversité et de l’Écotourisme")
    run.bold = True
    run.font.size = Pt(9)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("DEMANDE DE CONGÉ")
    r.bold = True
    r.font.size = Pt(18)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Document prêt pour transmission aux Ressources Humaines").italic = True

    doc.add_paragraph("")
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    fields = [
        ("Nom et prénom", data.get("full_name", "")),
        ("Poste", data.get("poste", "")),
        ("Type de congé", data.get("leave_type", "")),
        ("Motif", data.get("motif", "")),
        ("Date de début", data.get("start_date", "")),
        ("Date de fin", data.get("end_date", "")),
        ("Nombre de jours", str(data.get("nb_days", ""))),
        ("Personne assurant l’intérim", data.get("interim", "")),
        ("Contact pendant le congé", data.get("contact", "")),
        ("Responsable hiérarchique", data.get("manager", "")),
    ]
    for label, value in fields:
        row = table.add_row().cells
        row[0].text = label
        row[1].text = str(value or "")
        row[0].paragraphs[0].runs[0].bold = True
        for c in row:
            for p in c.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)

    doc.add_paragraph("")
    p = doc.add_paragraph()
    p.add_run("Déclaration du demandeur : ").bold = True
    p.add_run("Je sollicite l’autorisation de m’absenter pour la période indiquée ci-dessus et m’engage à assurer la transmission des informations nécessaires avant mon départ.")

    doc.add_paragraph("")
    sig = doc.add_table(rows=3, cols=3)
    sig.style = "Table Grid"
    headers = ["Demandeur", "Responsable hiérarchique", "Ressources Humaines"]
    for i, h in enumerate(headers):
        sig.cell(0, i).text = h
        sig.cell(0, i).paragraphs[0].runs[0].bold = True
    for i in range(3):
        sig.cell(1, i).text = "Date : ____ / ____ / ______"
        sig.cell(2, i).text = "Signature :"

    doc.add_paragraph("")
    foot = doc.add_paragraph("Fait à ______________________, le ____ / ____ / ______")
    foot.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()
